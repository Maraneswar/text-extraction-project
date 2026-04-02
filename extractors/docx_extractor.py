"""
DOCX text extraction using python-docx.
Extracts text preserving paragraph structure, tables, and document properties.
"""
import time
import os
from docx import Document
from models.schemas import ExtractionResult, DocumentMetadata


def extract_docx(file_path: str) -> ExtractionResult:
    """Extract text and metadata from a DOCX file."""
    start_time = time.time()

    try:
        doc = Document(file_path)

        # Extract paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Preserve heading structure
                if para.style and para.style.name.startswith("Heading"):
                    level = para.style.name.replace("Heading ", "").strip()
                    prefix = "#" * int(level) if level.isdigit() else "##"
                    paragraphs.append(f"{prefix} {text}")
                else:
                    paragraphs.append(text)

        # Extract tables
        tables_text = []
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(" | ".join(row_data))
            if table_data:
                tables_text.append(f"\n[Table {table_idx + 1}]\n" + "\n".join(table_data))

        # Combine all text
        full_text = "\n\n".join(paragraphs)
        if tables_text:
            full_text += "\n\n" + "\n".join(tables_text)

        # Extract metadata from core properties
        props = doc.core_properties
        metadata = DocumentMetadata(
            title=props.title or os.path.basename(file_path),
            author=props.author or "Unknown",
            creation_date=str(props.created) if props.created else "",
            modification_date=str(props.modified) if props.modified else "",
            page_count=None,  # DOCX doesn't expose page count easily
            word_count=len(full_text.split()) if full_text else 0,
            character_count=len(full_text),
            file_type="DOCX",
            extra={
                "category": props.category or "",
                "comments": props.comments or "",
                "last_modified_by": props.last_modified_by or "",
                "revision": props.revision,
                "subject": props.subject or "",
                "keywords": props.keywords or "",
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
            }
        )

        elapsed = (time.time() - start_time) * 1000

        if not full_text.strip():
            return ExtractionResult(
                raw_text="",
                metadata=metadata,
                success=False,
                error_message="No text content found in the DOCX file.",
                extraction_time_ms=elapsed,
            )

        return ExtractionResult(
            raw_text=full_text,
            metadata=metadata,
            success=True,
            extraction_time_ms=elapsed,
        )

    except Exception as e:
        elapsed = (time.time() - start_time) * 1000
        return ExtractionResult(
            raw_text="",
            metadata=DocumentMetadata(file_type="DOCX"),
            success=False,
            error_message=f"DOCX extraction failed: {str(e)}",
            extraction_time_ms=elapsed,
        )
